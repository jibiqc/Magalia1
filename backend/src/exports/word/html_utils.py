import bleach
import html
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from .settings import TEXT_COLOR, NORMAL_PT, DAY_PT

# Whitelist agreed
ALLOWED_TAGS = ["b", "strong", "i", "em", "ul", "ol", "li", "a", "p", "br"]
ALLOWED_ATTRS = {"a": ["href", "title"]}

def _apply_run_defaults(run):
    run.font.name = "Arial"
    run.font.size = Pt(NORMAL_PT)
    run.font.color.rgb = TEXT_COLOR
    run.italic = False  # no italics by default
    return run

def sanitize_html(html: str) -> str:
    return bleach.clean(html or "", tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)

def add_hyperlink(paragraph: Paragraph, url: str, text: str) -> None:
    """
    Create a clickable hyperlink in a paragraph.
    """
    # Build w:hyperlink with a relationship id on the part
    part = paragraph.part
    r_id = part.relate_to(url, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    # Create run
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Apply basic hyperlink style (blue + underline by Word default style)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    # Apply color and underline
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "002060")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def append_sanitized_html_to_docx(paragraph: Paragraph, html_text: str, allow_italic: bool = False) -> None:
    """
    Minimal HTML → Word:
    - <p>/<br>: new paragraphs only when needed (no duplicates)
    - <a>: clickable hyperlink runs
    - <ul>/<ol>/<li>: plain paragraphs with simple bullet/number prefix to avoid template style drift
    """
    safe = sanitize_html(html_text)
    import re

    def _emit_text_or_link(p: Paragraph, chunk: str, allow_italic: bool = False):
        # Gérer les balises <strong> et <b> pour le gras
        # D'abord, traiter les liens
        pos = 0
        for m in re.finditer(r'<a[^>]*href="([^"]+)"[^>]*>(.*?)</a>', chunk, flags=re.I):
            # Texte avant le lien (peut contenir <strong>)
            before = chunk[pos:m.start()]
            _add_text_with_formatting(p, before, allow_italic)
            # Lien
            url, text = m.group(1), bleach.clean(m.group(2), tags=[], strip=True)
            # Décoder les entités HTML
            if text:
                text = html.unescape(text)
            if url and text:
                add_hyperlink(p, url, text)
            pos = m.end()
        # Texte après le dernier lien (peut contenir <strong>)
        tail = chunk[pos:]
        _add_text_with_formatting(p, tail, allow_italic)
    
    def _add_text_with_formatting(p: Paragraph, text: str, allow_italic: bool = False):
        """Ajoute du texte en préservant le formatage <strong>/<b> pour le gras et <em>/<i> pour l'italique."""
        if not text:
            return
        
        # Fonction récursive pour traiter le texte avec formatage imbriqué
        def _process_text_segment(segment: str, is_bold: bool = False, is_italic: bool = False):
            """Traite un segment de texte avec formatage."""
            if not segment:
                return
            
            # Chercher toutes les balises (gras et italique)
            # Pattern pour trouver toutes les balises d'ouverture et de fermeture
            pattern = r'<(strong|b|em|i)>(.*?)</(strong|b|em|i)>'
            matches = list(re.finditer(pattern, segment, flags=re.I))
            
            if not matches:
                # Pas de formatage, ajouter le texte tel quel
                clean_text = bleach.clean(segment, tags=[], strip=True)
                if clean_text:
                    clean_text = html.unescape(clean_text)
                    if clean_text:
                        run = p.add_run(clean_text)
                        run.font.name = "Arial"
                        run.font.size = Pt(NORMAL_PT)
                        run.font.color.rgb = TEXT_COLOR
                        run.font.bold = is_bold
                        run.font.italic = is_italic or allow_italic
                return
            
            # Traiter la première balise trouvée
            first_match = matches[0]
            tag_name = first_match.group(1).lower()
            
            # Texte avant la première balise
            before = segment[:first_match.start()]
            _process_text_segment(before, is_bold, is_italic)
            
            # Déterminer le formatage de la balise
            new_bold = is_bold or (tag_name in ['strong', 'b'])
            new_italic = is_italic or (tag_name in ['em', 'i'])
            
            # Texte dans la balise (peut contenir d'autres balises imbriquées)
            inner_text = first_match.group(2)
            _process_text_segment(inner_text, new_bold, new_italic)
            
            # Texte après la première balise
            after = segment[first_match.end():]
            _process_text_segment(after, is_bold, is_italic)
        
        # Traiter le texte complet
        _process_text_segment(text)

    # Normalize paragraphs and lists
    # Convert </p> to \n markers, strip <p>, handle <br> as line breaks
    text = safe.replace("</p>", "\n").replace("<p>", "")
    # Lists: split blocks for <ul> and <ol>
    # We render each <li> as its own Normal paragraph with a simple prefix.
    # Unordered
    def _render_list(par: Paragraph, block: str, ordered: bool, allow_italic: bool = False):
        items = re.findall(r"<li>(.*?)</li>", block, flags=re.I | re.S)
        for i, raw in enumerate(items, 1):
            newp = par._parent.add_paragraph()
            newp.style = par.style
            prefix = f"{i}. " if ordered else "• "
            run = newp.add_run(prefix)
            if not allow_italic:
                run = _apply_run_defaults(run)
            _emit_text_or_link(newp, raw, allow_italic=allow_italic)

    # Process lists first
    def _strip_lists(s: str, par: Paragraph, allow_italic: bool = False) -> str:
        # Ordered lists
        for m in re.finditer(r"<ol>(.*?)</ol>", s, flags=re.I | re.S):
            _render_list(par, m.group(1), ordered=True, allow_italic=allow_italic)
        s = re.sub(r"<ol>.*?</ol>", "", s, flags=re.I | re.S)
        # Unordered lists
        for m in re.finditer(r"<ul>(.*?)</ul>", s, flags=re.I | re.S):
            _render_list(par, m.group(1), ordered=False, allow_italic=allow_italic)
        s = re.sub(r"<ul>.*?</ul>", "", s, flags=re.I | re.S)
        return s

    text = _strip_lists(text, paragraph, allow_italic=allow_italic)
    # Split leftover by explicit paragraph breaks
    for idx, blk in enumerate([b for b in text.split("\n") if b.strip()]):
        if idx == 0:
            _emit_text_or_link(paragraph, blk, allow_italic=allow_italic)
        else:
            newp = paragraph._parent.add_paragraph()
            newp.style = paragraph.style
            _emit_text_or_link(newp, blk, allow_italic=allow_italic)

