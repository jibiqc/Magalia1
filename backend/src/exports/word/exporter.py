from __future__ import annotations
from io import BytesIO
from typing import Optional, List
import re
import html
import calendar
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from sqlalchemy.orm import Session

from .settings import (
    WORD_TEMPLATE_PATH, HERO_W_CM, HERO_H_CM,
    TEXT_COLOR, TITLE_PT, DAY_PT, DATE_PT, NORMAL_PT,
    SP_AFTER_NORMAL_PT, SP_BEFORE_ALL_PT, SP_AFTER_ALL_PT
)
from .image_utils import fetch_image_bytes, contain_resize_to_width_cm, stitch_side_by_side_to_cm
from .html_utils import append_sanitized_html_to_docx, add_hyperlink
from ...models_quote import Quote

# --- Date formatting identical to UI, English locale, ordinals ---
ORD = {1: "st", 2: "nd", 3: "rd"}

def _ordinal(n: int) -> str:
    return f"{n}{ORD.get(0 if 10 <= n % 100 <= 20 else n % 10, 'th')}"

def fmt_ui_date(d_iso: str) -> str:
    # "Tuesday, November 11th 2025 :"
    dt = datetime.fromisoformat(d_iso)
    wd = calendar.day_name[dt.weekday()]
    month = calendar.month_name[dt.month]
    return f"{wd}, {month} {_ordinal(dt.day)} {dt.year} :"

def _is_first_of_dest_block(days, idx):
    cur = (_get_attr(days[idx], "destination") or "").strip()
    if not cur:
        return False
    if idx == 0:
        return True
    prev = (_get_attr(days[idx-1], "destination") or "").strip()
    return cur != prev

def _count_nights_from_index(days, idx):
    cur = (_get_attr(days[idx], "destination") or "").strip()
    n = 1
    j = idx + 1
    while j < len(days) and (_get_attr(days[j], "destination") or "").strip() == cur:
        n += 1
        j += 1
    return n

# --- Paragraph factories ---
def _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT, line_single=True):
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)
    if line_single:
        p.paragraph_format.line_spacing = 1.0

def _add_title(doc: Document, text: str):
    # Titre global : Arial 28, gras, couleur #002060, centré
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(TITLE_PT)
    r.font.bold = True
    r.font.color.rgb = TEXT_COLOR
    r.italic = False
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT)
    # ← ajout : insérer une ligne vide après le titre
    doc.add_paragraph()
    return p

def _add_date_line(doc: Document, text: str):
    # Date: Arial 10, underlined, NOT bold, color; after 0; no italics
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(DATE_PT)
    r.font.bold = False
    r.underline = True
    r.font.color.rgb = TEXT_COLOR
    r.italic = False
    _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT)
    return p

def _add_title_with_ordinals(p: Paragraph, text: str):
    """Ajoute un titre avec conversion des ordinaux (1st, 2nd, 3rd, etc.) en exposants."""
    if not text or not isinstance(text, str):
        return
    
    # Pattern pour détecter les ordinaux: 1st, 2nd, 3rd, 4th, etc.
    ordinal_pattern = re.compile(r'(\d+)(st|nd|rd|th)', re.IGNORECASE)
    
    parts = []
    last_index = 0
    
    for match in ordinal_pattern.finditer(text):
        # Texte avant l'ordinal
        if match.start() > last_index:
            parts.append(('normal', text[last_index:match.start()]))
        
        # Nombre et suffixe ordinal
        parts.append(('number', match.group(1)))
        parts.append(('superscript', match.group(2)))
        
        last_index = match.end()
    
    # Texte restant après le dernier ordinal
    if last_index < len(text):
        parts.append(('normal', text[last_index:]))
    
    # Si aucun ordinal trouvé, ajouter tout le texte normalement
    if not parts:
        parts.append(('normal', text))
    
    # Créer les runs Word
    for part_type, part_text in parts:
        if not part_text:
            continue
        
        run = p.add_run(part_text)
        run.font.name = "Arial"
        run.font.size = Pt(DAY_PT)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        run.font.italic = False
        
        if part_type == 'superscript':
            run.font.superscript = True

def _add_day_title(doc: Document, text: str):
    # Day/service titles: Arial 10, bold, color; after 0
    p = doc.add_paragraph()
    _add_title_with_ordinals(p, text)
    _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT)
    return p

def _add_normal(doc: Document, text: str):
    # Normal: Arial 10, color; after 10 pt; no italics
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(NORMAL_PT)
    r.font.color.rgb = TEXT_COLOR
    r.italic = False
    _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_NORMAL_PT)
    return p

def _blank_line(doc: Document):
    p = doc.add_paragraph()
    _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT)
    return p

def _add_title_with_html_formatting(p: Paragraph, html_text: str):
    """Ajoute un titre avec formatage HTML (gras pour <strong>/<b>)."""
    import bleach
    
    try:
        # S'assurer que html_text est une chaîne
        if not html_text or not isinstance(html_text, str):
            return
        
        # Vérifier s'il y a des balises HTML
        has_html_tags = bool(re.search(r'<(strong|b)>', html_text, flags=re.I))
        
        if has_html_tags:
            # Extraire les segments avec/sans gras
            pos = 0
            for m in re.finditer(r'<(strong|b)>(.*?)</(strong|b)>', html_text, flags=re.I):
                # Texte avant le gras
                before = bleach.clean(html_text[pos:m.start()], tags=[], strip=True)
                # Décoder les entités HTML comme &amp; en &
                if before:
                    before = html.unescape(before)
                if before:
                    run = p.add_run(before)
                    run.font.name = "Arial"
                    run.font.size = Pt(DAY_PT)
                    run.font.bold = True  # Titre en gras par défaut
                    run.font.color.rgb = TEXT_COLOR
                    run.font.italic = False
                # Texte en gras (déjà en gras, donc on garde)
                bold_text = bleach.clean(m.group(2), tags=[], strip=True)
                # Décoder les entités HTML
                if bold_text:
                    bold_text = html.unescape(bold_text)
                if bold_text:
                    run = p.add_run(bold_text)
                    run.font.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(DAY_PT)
                    run.font.color.rgb = TEXT_COLOR
                    run.font.italic = False
                pos = m.end()
            # Texte restant après le dernier gras
            remaining = bleach.clean(html_text[pos:], tags=[], strip=True)
            # Décoder les entités HTML
            if remaining:
                remaining = html.unescape(remaining)
            if remaining:
                run = p.add_run(remaining)
                run.font.name = "Arial"
                run.font.size = Pt(DAY_PT)
                run.font.bold = True  # Titre en gras par défaut
                run.font.color.rgb = TEXT_COLOR
                run.font.italic = False
        else:
            # Si pas de HTML, ajouter tout le texte en gras
            clean_text = bleach.clean(html_text, tags=[], strip=True)
            # Décoder les entités HTML
            if clean_text:
                clean_text = html.unescape(clean_text)
            if clean_text:
                run = p.add_run(clean_text)
                run.font.name = "Arial"
                run.font.size = Pt(DAY_PT)
                run.font.bold = True
                run.font.color.rgb = TEXT_COLOR
                run.font.italic = False
    except Exception as e:
        # En cas d'erreur, ajouter le texte brut sans formatage HTML
        print(f"WARNING: Error parsing HTML in title: {e}, using plain text")
        try:
            clean_text = bleach.clean(str(html_text), tags=[], strip=True) if html_text else ""
            if clean_text:
                clean_text = html.unescape(clean_text)
            if clean_text:
                run = p.add_run(clean_text)
                run.font.name = "Arial"
                run.font.size = Pt(DAY_PT)
                run.font.bold = True
                run.font.color.rgb = TEXT_COLOR
                run.font.italic = False
        except Exception:
            # Dernier recours : texte brut
            if html_text:
                run = p.add_run(str(html_text))
                run.font.name = "Arial"
                run.font.size = Pt(DAY_PT)
                run.font.bold = True
                run.font.color.rgb = TEXT_COLOR
                run.font.italic = False

# --- Helper pour accéder aux attributs (dict ou objet) ---
def _get_attr(obj, key, default=None):
    """Helper pour accéder à un attribut/clé d'un objet ou dictionnaire."""
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)

# --- Service title composition (identical to ServiceCard.jsx) ---
def _norm_yes(v) -> bool:
    s = str(v or "").lower()
    return s == "1" or s == "yes" or "breakfast" in s

def _norm_stars(v) -> str:
    s = str(v or "").strip()
    if not s:
        return ""
    try:
        n = round(float("".join(ch for ch in s if (ch.isdigit() or ch == "."))))
    except Exception:
        return ""
    n = max(1, min(5, n))
    return " " + "★" * n

def _infer_room(t: str) -> str:
    t = (t or "")
    t = re.sub(r"\bat\b.*$", "", t, flags=re.I)
    t = re.sub(r"^(hotel\s*)?room\b[:\-]?\s*", "", t, flags=re.I).strip()
    return t or "Room"

def _fmt_ampm(hhmm: str) -> str:
    if not hhmm:
        return ""
    parts = hhmm.split(":")
    h = int(parts[0] or "0")
    m = (parts[1] if len(parts) > 1 else "00").zfill(2)
    suf = "pm" if h >= 12 else "am"
    h = (h % 12) or 12
    return f"{h}:{m} {suf}"

def _pretty_duration(duration_str: str) -> str:
    """Formate une durée comme dans ServiceCard (prettyDur)."""
    if not duration_str:
        return ""
    duration_str = str(duration_str).strip()
    if not duration_str:
        return ""
    
    # Pattern: 3h30 ou 3:30 ou "3 30"
    m = re.match(r'^(\d+)\s*(?:h|:|\s)\s*(\d{1,2})$', duration_str, re.I)
    if m:
        total_mins = int(m.group(1)) * 60 + int(m.group(2))
        h = total_mins // 60
        m = total_mins % 60
        if h and m:
            return f"{h}h {m}m"
        if h:
            return f"{h}h"
        return f"{m}m"
    
    # Pattern: 3h
    m = re.match(r'^(\d+)\s*h$', duration_str, re.I)
    if m:
        h = int(m.group(1))
        return f"{h}h"
    
    # Pattern: 90m / 90min
    m = re.match(r'^(\d+)\s*m(?:in)?$', duration_str, re.I)
    if m:
        total_mins = int(m.group(1))
        h = total_mins // 60
        m = total_mins % 60
        if h and m:
            return f"{h}h {m}m"
        if h:
            return f"{h}h"
        return f"{m}m"
    
    # Pattern: plain minutes "150"
    if re.match(r'^\d+$', duration_str):
        total_mins = int(duration_str)
        h = total_mins // 60
        m = total_mins % 60
        if h and m:
            return f"{h}h {m}m"
        if h:
            return f"{h}h"
        return f"{m}m"
    
    # Sinon, retourner tel quel
    return duration_str

def _fmt_service_title(line):
    cat = (_get_attr(line, "category") or "").strip()
    title = _get_attr(line, "title") or cat or "Service"
    rj = _get_attr(line, "raw_json", {}) or {}

    # HOTEL
    if cat == "Hotel":
        isCatalog = bool(rj.get("catalog_id"))
        if isCatalog:
            room = (rj.get("room_type") or "").strip()
            company = (rj.get("hotel_name") or _get_attr(line, "provider_name") or _get_attr(line, "supplier_name") or "").strip()
            stars = _norm_stars(rj.get("hotel_stars"))
        else:
            room = _infer_room(title)
            company = _get_attr(line, "provider_name") or _get_attr(line, "supplier_name") or ""
            stars = _norm_stars(rj.get("hotel_stars"))
        bf_json = rj.get("breakfast")
        if bf_json is True:
            bf = "breakfast & VAT included"
        elif bf_json is False:
            bf = "VAT included"
        else:
            bf = "breakfast & VAT included" if _norm_yes(rj.get("meal_1")) else "VAT included"
        eci = ", early check-in guaranteed" if rj.get("early_check_in") else ""
        title = f"{room}{eci}, {bf} at {company}{stars}"

    # ACTIVITÉS
    elif cat in ("Activity", "Small Group", "Private"):
        st = rj.get("start_time") or ""
        if st:
            # Utiliser _fmt_ampm si le format contient ":"
            st_formatted = _fmt_ampm(st) if ":" in st else st
            title = f"{title} at {st_formatted}"

    # TRANSPORTS
    elif cat == "Train":
        # Reconstruire le titre comme dans ServiceCard avec "train" en minuscule
        data = rj
        class_type = data.get("class_type") or "First Class"
        from_loc = data.get("from") or "?"
        to_loc = data.get("to") or "?"
        suffix = ""
        choice = data.get("seat_res_choice")
        if choice == "with":
            suffix = " with seat reservations"
        elif choice == "without":
            suffix = " without seat reservations (open seating)"
        elif choice == "none":
            suffix = ""
        elif data.get("seat_res") is True:
            suffix = " with seat reservations"
        elif data.get("seat_res") is False:
            suffix = " without seat reservations (open seating)"
        title = f"{class_type} train from {from_loc} to {to_loc}{suffix}"
    elif cat == "Ferry":
        # Reconstruire le titre comme dans ServiceCard
        data = rj
        cls = (data.get("class_type") or "").strip()
        cls_prefix = f"{cls} " if cls else ""
        from_loc = data.get("from") or "?"
        to_loc = data.get("to") or "?"
        suffix = ""
        choice = data.get("seat_res_choice")
        if choice == "with":
            suffix = " with seat reservations"
        elif choice == "without":
            suffix = " without seat reservations (open seating)"
        elif choice == "none":
            suffix = ""
        elif data.get("seat_res") is True:
            suffix = " with seat reservations"
        elif data.get("seat_res") is False:
            suffix = " without seat reservations (open seating)"
        title = f"{cls_prefix}Ferry from {from_loc} to {to_loc}{suffix}"
    elif cat == "Flight":
        # Reconstruire le titre comme dans ServiceCard avec class_of_service
        data = rj
        airline = data.get("airline") or "Airline"
        cls = (data.get("class_of_service") or "").strip()
        cls_prefix = f"{cls} " if cls else ""
        from_loc = data.get("from") or "?"
        to_loc = data.get("to") or "?"
        suffix = ""
        choice = data.get("seat_res_opt")
        if choice == "with":
            suffix = " with seat reservations"
        elif data.get("seat_res") is True or data.get("with_seats") is True:
            suffix = " with seat reservations"
        title = f"{cls_prefix}{airline} flight from {from_loc} to {to_loc}{suffix}"
    elif cat == "Private Transfer":
        # Afficher tel quel, pas de normalisation destructrice
        title = _get_attr(line, "title") or title
    # CAR RENTAL - construire le titre comme dans ServiceCard
    elif cat == "Car Rental":
        data = rj  # raw_json contient les données du modal
        loc = (data.get("pickup_loc") or "?").strip()
        at_air = ""
        if data.get("pickup_airport"):
            at_air = " " + str(data.get("pickup_airport")).strip()
        vehicle = (data.get("vehicle_type") or "").strip()
        tx = ""
        transmission = data.get("transmission") or ""
        if transmission and transmission != "Do not precise":
            tx = f"{transmission.lower()}, "
        mileage = (data.get("mileage") or "").strip()
        ins = (data.get("insurance") or "").strip()
        # Construire le titre comme dans ServiceCard
        title = f"Pick up car in {loc}{at_air}, {vehicle} {tx}{mileage} {ins}"
        # Normaliser les espaces multiples en un seul espace (comme dans ServiceCard)
        title = re.sub(r'\s+', ' ', title).strip()
        # Si le titre est vide ou juste "Pick up car in ?, ,  ", utiliser le titre original
        if not title or title == "Pick up car in ?, ,  ":
            title = _get_attr(line, "title") or title

    # NEW HOTEL
    elif cat == "New Hotel":
        # Les données sont dans raw_json (ou data pour les lignes locales)
        data = rj  # raw_json contient les données du modal
        stars_val = data.get("stars") or ""
        star_count = 0
        if stars_val and stars_val != "NA":
            try:
                star_count = int(float(str(stars_val).strip()))
                star_count = max(1, min(5, star_count))
            except (ValueError, TypeError):
                pass
        stars = f" {'★' * star_count}" if star_count > 0 else ""
        
        room = (data.get("room_type") or "").strip()
        opts = []
        if data.get("early_checkin"):
            opts.append("early check-in guaranteed")
        opts.append("breakfast & VAT taxes included" if data.get("breakfast") else "VAT taxes included")
        hotel_name = (data.get("hotel_name") or "Hotel").strip()
        
        title_parts = []
        if room:
            title_parts.append(room)
        if opts:
            title_parts.append(", ".join(opts))
        title_parts.append(f"at {hotel_name}{stars}")
        title = ", ".join(title_parts)

    # NEW SERVICE
    elif cat == "New Service":
        # Les données sont dans raw_json (ou data pour les lignes locales)
        data = rj  # raw_json contient les données du modal
        service_title = (data.get("title") or "Service").strip()
        start_time = data.get("start_time") or ""
        if start_time:
            st_formatted = _fmt_ampm(start_time) if ":" in start_time else start_time
            title = f"{service_title} at {st_formatted}"
        else:
            title = service_title

    return title

def _is_trip_info(line) -> bool:
    return (_get_attr(line, "category") or "").strip() == "Trip info"

def _is_transport_service(category: str) -> bool:
    """Vérifie si une catégorie est un service de transport."""
    transport_categories = ("Private Transfer", "Train", "Ferry", "Flight")
    return category in transport_categories

def _append_hotel_url_if_any(doc, line):
    raw_json = _get_attr(line, "raw_json", {}) or {}
    url = (raw_json.get("hotel_url") or "").strip()
    if url:
        p = doc.add_paragraph()
        r = p.add_run(url)
        r.font.name = "Arial"
        r.font.size = Pt(NORMAL_PT)
        r.font.color.rgb = TEXT_COLOR
        r.underline = True
        r.italic = False
        _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_NORMAL_PT)

def _insert_two_heroes(doc: Document, url1: str, url2: str, add_blank_after: bool = True):
    if not url1 or not url2:
        return None, None
    # Pas d'espace avant la photo
    try:
        jpeg = stitch_side_by_side_to_cm(url1, url2, width_cm=HERO_W_CM, height_cm=HERO_H_CM)
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(BytesIO(jpeg), width=Cm(HERO_W_CM), height=Cm(HERO_H_CM))
        _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT)
        # Ligne vide après la photo (sauf si add_blank_after est False, pour garder la photo avec la date)
        blank_para = None
        if add_blank_after:
            blank_para = _blank_line(doc)
        return p, blank_para
    except Exception:
        pass  # Skip if image fetch fails
    return None, None

def _render_service(doc: Document, line, usable_width_cm: float, skip_blank_line_after: bool = False):
    """Rend un service dans le document Word avec gestion d'erreur robuste."""
    try:
        cat = (_get_attr(line, "category") or "").strip()
        rj = _get_attr(line, "raw_json", {}) or {}
        
        # Mémoriser l'index du premier paragraphe du service pour gérer les sauts de page
        # Cela permet de garder tous les éléments du service (titre, sous-titre, description, URL) ensemble
        service_start_para_idx = len(doc.paragraphs)
        
        # 1) Titre - toujours affiché
        title = _fmt_service_title(line)
        if title:
            # Pour Trip info, titre en italique (pas de gras)
            if cat == "Trip info":
                p = doc.add_paragraph()
                # Utiliser _add_title_with_ordinals mais avec italique au lieu de gras
                if title:
                    ordinal_pattern = re.compile(r'(\d+)(st|nd|rd|th)', re.IGNORECASE)
                    parts = []
                    last_index = 0
                    
                    for match in ordinal_pattern.finditer(title):
                        if match.start() > last_index:
                            parts.append(('normal', title[last_index:match.start()]))
                        parts.append(('number', match.group(1)))
                        parts.append(('superscript', match.group(2)))
                        last_index = match.end()
                    
                    if last_index < len(title):
                        parts.append(('normal', title[last_index:]))
                    
                    if not parts:
                        parts.append(('normal', title))
                    
                    for part_type, part_text in parts:
                        if not part_text:
                            continue
                        r = p.add_run(part_text)
                        r.font.name = "Arial"
                        r.font.size = Pt(DAY_PT)
                        r.font.bold = False
                        r.font.italic = True
                        r.font.color.rgb = TEXT_COLOR
                        if part_type == 'superscript':
                            r.font.superscript = True
                _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT)
            # Pour Car Rental, parser le HTML pour préserver le formatage (gras pour SUV, CDW, etc.)
            elif cat == "Car Rental":
                p = doc.add_paragraph()
                # Le titre peut contenir du HTML avec <strong> pour SUV, CDW, etc.
                # Utiliser une fonction spéciale pour préserver le formatage
                _add_title_with_html_formatting(p, title)
                _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_ALL_PT)
            else:
                _add_day_title(doc, title)
        added_any = False

        # 2) Sous-titre pour Flight, Train, Ferry (heures de départ/arrivée)
        if cat == "Flight":
            dep_time = rj.get("dep_time") or ""
            arr_time = rj.get("arr_time") or ""
            if dep_time or arr_time:
                subtitle_parts = []
                if dep_time:
                    subtitle_parts.append(f"Departure at {_fmt_ampm(dep_time)}")
                if arr_time:
                    subtitle_parts.append(f"arrival at {_fmt_ampm(arr_time)}")
                if subtitle_parts:
                    subtitle = f"{'; '.join(subtitle_parts)} – Schedule subject to change"
                    _add_normal(doc, subtitle)
                    added_any = True
    
        elif cat == "Train":
            dep_time = rj.get("dep_time") or ""
            arr_time = rj.get("arr_time") or ""
            if dep_time or arr_time:
                subtitle_parts = []
                if dep_time:
                    subtitle_parts.append(f"Departure at {_fmt_ampm(dep_time)}")
                if arr_time:
                    subtitle_parts.append(f"arrival at {_fmt_ampm(arr_time)}")
                if subtitle_parts:
                    subtitle = f"{'; '.join(subtitle_parts)} – Schedule subject to change"
                    _add_normal(doc, subtitle)
                    added_any = True
        
        elif cat == "Ferry":
            dep_time = rj.get("dep_time") or ""
            arr_time = rj.get("arr_time") or ""
            if dep_time or arr_time:
                subtitle_parts = []
                if dep_time:
                    subtitle_parts.append(f"Departure {_fmt_ampm(dep_time)}")
                if arr_time:
                    subtitle_parts.append(f"Arrival {_fmt_ampm(arr_time)}")
                if subtitle_parts:
                    subtitle = f"{'; '.join(subtitle_parts)} – Schedule subject to change"
                    _add_normal(doc, subtitle)
                    added_any = True

        # Description (corps) - affichée après les sous-titres
        fields = rj.get("fields", {}) or {}
        desc = fields.get("full_description") or rj.get("description") or rj.get("note") or _get_attr(line, "description", "") or ""
        
        # Pour les activités, ajouter la duration au début de la description si elle existe
        # (pour l'export Excel, même si elle est déjà dans le sous-titre)
        if cat in ("Activity", "Small Group", "Private"):
            # Chercher la duration dans raw_json ou fields
            duration = rj.get("duration") or fields.get("activity_duration") or ""
            if duration and str(duration).strip():
                duration_formatted = _pretty_duration(str(duration))
                if duration_formatted:
                    # Si la description existe, ajouter la duration au début avec un saut de ligne
                    if desc:
                        desc = f"Duration: {duration_formatted}\n{desc}"
                    else:
                        # Si pas de description, créer une description avec juste la duration
                        desc = f"Duration: {duration_formatted}"
        
        # Pour Car Rental, construire la description comme dans ServiceCard (feeLine + desc + licenceLine)
        if cat == "Car Rental":
            data = rj  # raw_json contient les données du modal
            fee_line = ""
            one_way_fee = data.get("one_way_fee")
            if one_way_fee and float(str(one_way_fee).replace(",", ".") or "0") > 0:
                try:
                    fee_value = float(str(one_way_fee).replace(",", "."))
                    fee_line = f"Estimate One Way Fee: ${int(fee_value)} – to be paid locally"
                except (ValueError, TypeError):
                    pass
            
            car_desc = data.get("description") or data.get("notes") or desc or ""
            
            # Fonction helper pour normaliser et comparer les textes
            def _normalize_for_comparison(text):
                """Normalise le texte pour la comparaison : supprime HTML, normalise espaces, en minuscules"""
                if not text:
                    return ""
                # Supprimer HTML
                text_no_html = re.sub(r'<[^>]+>', '', text)
                # Normaliser espaces
                text_normalized = re.sub(r'\s+', ' ', text_no_html.strip())
                return text_normalized.lower()
            
            # Éviter la duplication : si car_desc contient le titre (ou est identique au titre), ne pas l'inclure
            # Le titre est déjà affiché séparément, donc on ne veut pas le répéter dans la description
            if title and car_desc:
                title_norm = _normalize_for_comparison(title)
                car_desc_norm = _normalize_for_comparison(car_desc)
                
                # Si car_desc est identique au titre ou contient le titre, le vider
                if car_desc_norm == title_norm or title_norm in car_desc_norm:
                    car_desc = ""
                # Vérifier aussi si le titre commence par "pick up car" et si car_desc commence de la même manière
                elif title_norm.startswith("pick up car") and car_desc_norm.startswith("pick up car"):
                    # Si les deux commencent par "pick up car", comparer les premiers mots (jusqu'à 10 mots pour être sûr)
                    title_words = title_norm.split()[:10]  # Premiers 10 mots
                    car_desc_words = car_desc_norm.split()[:10]
                    # Comparer si les mots sont identiques
                    if len(title_words) > 0 and title_words == car_desc_words[:len(title_words)]:
                        car_desc = ""
            
            licence_line = ""
            if data.get("intl_driver_license"):
                licence_line = "An international driver's license is mandatory to pick up the car. A physical hard copy is required, as digital copies are not accepted locally. Please note that it may take up to 15 days to obtain the license."
            
            # Construire la description complète comme dans ServiceCard
            desc_parts = [fee_line, car_desc, licence_line]
            desc = "\n\n".join([p for p in desc_parts if p.strip()])
    
        # Pour New Service, la description est déjà dans le sous-titre, donc on ne la répète pas
        # Pour New Hotel, utiliser description depuis raw_json
        elif cat == "New Hotel":
            desc = rj.get("description") or ""
        elif cat == "New Service":
            # La description est déjà dans le sous-titre, donc on ne l'affiche pas ici
            desc = ""
        
        # Pour ces catégories spécifiques, on s'assure d'exporter le texte même s'il est dans le titre
        # MAIS: pour Flight/Train/Ferry, si on a déjà affiché le sous-titre avec les heures, on n'utilise PAS le titre comme description
        has_subtitle = cat in ("Flight", "Train", "Ferry") and (rj.get("dep_time") or rj.get("arr_time"))
        
        if cat in ("Flight", "Train", "Ferry", "Trip info"):
            # Si pas de description et qu'on n'a pas de sous-titre, on peut utiliser le titre comme description (fallback)
            if not desc and title and not has_subtitle:
                desc = title

        # Mémoriser l'index du premier paragraphe de la description (si elle existe)
        desc_start_para_idx = None
        desc_end_para_idx = None
        
        if desc:
            # Mémoriser l'index du premier paragraphe de la description
            desc_start_para_idx = len(doc.paragraphs)
            
            # Pour Trip info, description en italique
            para_count_before = len(doc.paragraphs)
            p_desc = doc.add_paragraph()
            if cat == "Trip info":
                # Description en italique (pas de gras) et en bleu
                append_sanitized_html_to_docx(p_desc, desc, allow_italic=True)
                # Forcer italique, pas de gras, couleur bleue, police Arial taille 10 pour tous les runs
                para_count_after = len(doc.paragraphs)
                for i in range(para_count_before, para_count_after):
                    for run in doc.paragraphs[i].runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(NORMAL_PT)  # NORMAL_PT = 10
                        run.font.italic = True
                        run.font.bold = False
                        run.font.color.rgb = TEXT_COLOR  # Forcer la couleur bleue
                    _set_par_spacing(doc.paragraphs[i], before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_NORMAL_PT)
            else:
                # Autres catégories : italique autorisé seulement pour Trip info (mais on est déjà dans le else)
                append_sanitized_html_to_docx(p_desc, desc, allow_italic=False)
                # Apply spacing to all generated paragraphs
                para_count_after = len(doc.paragraphs)
                for i in range(para_count_before, para_count_after):
                    _set_par_spacing(doc.paragraphs[i], before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_NORMAL_PT)
            added_any = True
            
            # Mémoriser l'index du dernier paragraphe de la description
            desc_end_para_idx = len(doc.paragraphs) - 1

        # 3) Sous-titre pour Activity (catalogue) avec duration
        if cat in ("Activity", "Small Group", "Private"):
            is_catalog_activity = bool(rj.get("catalog_id"))
            if is_catalog_activity:
                duration = rj.get("duration") or ""
                subtitle_parts = []
                if duration and str(duration).strip():
                    duration_formatted = _pretty_duration(str(duration))
                    if duration_formatted:
                        subtitle_parts.append(f"Duration: {duration_formatted}")
                
                # Description dans le sous-titre (comme dans ServiceCard)
                desc_for_subtitle = rj.get("description") or ""
                if desc_for_subtitle:
                    subtitle_parts.append(desc_for_subtitle)
                
                if subtitle_parts:
                    subtitle = "\n".join(subtitle_parts)
                    _add_normal(doc, subtitle)
                    added_any = True
            else:
                # Pour les activités non-catalogue, calculer la durée depuis start_time/end_time
                start_time = rj.get("start_time") or ""
                end_time = rj.get("end_time") or ""
                duration = rj.get("duration") or ""
                subtitle_parts = []
                
                # Calculer la durée si start_time et end_time sont disponibles
                if start_time and end_time and ":" in start_time and ":" in end_time:
                    try:
                        start_parts = start_time.split(":")
                        end_parts = end_time.split(":")
                        start_mins = int(start_parts[0]) * 60 + int(start_parts[1] if len(start_parts) > 1 else 0)
                        end_mins = int(end_parts[0]) * 60 + int(end_parts[1] if len(end_parts) > 1 else 0)
                        diff_mins = ((end_mins - start_mins + 1440) % 1440)
                        if diff_mins > 0:
                            hours = diff_mins // 60
                            mins = diff_mins % 60
                            if mins > 0:
                                duration_str = f"{hours}h{mins:02d}"
                            else:
                                duration_str = f"{hours}h"
                            subtitle_parts.append(f"Duration: {duration_str}")
                        elif diff_mins == -1:
                            subtitle_parts.append("⚠ End time before start time")
                    except (ValueError, IndexError):
                        pass
                
                # Sinon utiliser le champ duration s'il existe
                if not subtitle_parts and duration:
                    duration_formatted = _pretty_duration(str(duration))
                    if duration_formatted:
                        subtitle_parts.append(f"Duration: {duration_formatted}")
                
                # Description dans le sous-titre
                desc_for_subtitle = rj.get("description") or ""
                if desc_for_subtitle:
                    subtitle_parts.append(desc_for_subtitle)
                
                if subtitle_parts:
                    subtitle = "\n".join(subtitle_parts)
                    _add_normal(doc, subtitle)
                    added_any = True

        # 4) Sous-titre pour New Service (duration et description)
        elif cat == "New Service":
            data = rj  # raw_json contient les données du modal
            start_time = data.get("start_time") or ""
            end_time = data.get("end_time") or ""
            duration = data.get("duration") or ""
            
            subtitle_parts = []
            # Calculer la durée si start_time et end_time sont disponibles
            if start_time and end_time and ":" in start_time and ":" in end_time:
                try:
                    start_parts = start_time.split(":")
                    end_parts = end_time.split(":")
                    start_mins = int(start_parts[0]) * 60 + int(start_parts[1] if len(start_parts) > 1 else 0)
                    end_mins = int(end_parts[0]) * 60 + int(end_parts[1] if len(end_parts) > 1 else 0)
                    diff_mins = ((end_mins - start_mins + 1440) % 1440)
                    if diff_mins > 0:
                        hours = diff_mins // 60
                        mins = diff_mins % 60
                        if mins > 0:
                            duration_str = f"{hours}h{mins:02d}"
                        else:
                            duration_str = f"{hours}h"
                        subtitle_parts.append(f"Duration: {duration_str}")
                except (ValueError, IndexError):
                    pass
            
            # Sinon utiliser le champ duration s'il existe
            if not subtitle_parts and duration:
                duration_str = str(duration).strip()
                if duration_str:
                    subtitle_parts.append(f"Duration: {duration_str}")
            
            # Description dans le sous-titre (comme dans ServiceCard)
            desc_for_subtitle = data.get("description") or ""
            if desc_for_subtitle:
                subtitle_parts.append(desc_for_subtitle)
            
            if subtitle_parts:
                subtitle = "\n".join(subtitle_parts)
                _add_normal(doc, subtitle)
                added_any = True

        # 5) Lien hôtel si catégorie Hotel ou New Hotel
        if cat == "Hotel":
            _append_hotel_url_if_any(doc, line)
            added_any = True
        elif cat == "New Hotel":
            # Pour New Hotel, l'URL est dans raw_json.hotel_url
            data = rj
            url = (data.get("hotel_url") or "").strip()
            if url:
                p = doc.add_paragraph()
                r = p.add_run(url)
                r.font.name = "Arial"
                r.font.size = Pt(NORMAL_PT)
                r.font.color.rgb = TEXT_COLOR
                r.underline = True
                r.italic = False
                _set_par_spacing(p, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_NORMAL_PT)
                added_any = True

        # Mémoriser l'index du dernier paragraphe du service (avant le _blank_line)
        # C'est le dernier paragraphe qui fait partie du service (titre, sous-titre, description, URL)
        service_end_para_idx = len(doc.paragraphs) - 1
        
        # 4) Saut de ligne après le dernier élément du service (sauf si skip_blank_line_after est True)
        if not skip_blank_line_after:
            _blank_line(doc)
        
        # Appliquer keep_with_next et keep_together de manière sélective
        # Pour les activités et hôtels, permettre les sauts de page entre les paragraphes de description
        # pour éviter trop d'espace vide en bas de page
        if service_start_para_idx <= service_end_para_idx and service_end_para_idx < len(doc.paragraphs):
            # Déterminer si on doit permettre les sauts de page dans la description
            allow_page_breaks_in_desc = cat in ("Activity", "Small Group", "Private", "Hotel", "New Hotel")
            
            for para_idx in range(service_start_para_idx, service_end_para_idx + 1):
                para = doc.paragraphs[para_idx]
                
                # Si c'est un paragraphe de description pour activités/hôtels, permettre les sauts de page
                if (allow_page_breaks_in_desc and 
                    desc_start_para_idx is not None and 
                    desc_end_para_idx is not None and
                    desc_start_para_idx <= para_idx <= desc_end_para_idx):
                    # Pour les paragraphes de description (sauf le premier), ne pas forcer keep_with_next
                    # Cela permet à Word de faire des sauts de page entre les paragraphes si nécessaire
                    if para_idx == desc_start_para_idx:
                        # Le premier paragraphe de description reste avec le sous-titre (si existe)
                        para.paragraph_format.keep_with_next = True
                    else:
                        # Les autres paragraphes de description peuvent être séparés par un saut de page
                        para.paragraph_format.keep_with_next = False
                    # Toujours empêcher la division d'un paragraphe sur plusieurs pages
                    para.paragraph_format.keep_together = True
                else:
                    # Pour le titre, sous-titre, URL et autres éléments : garder ensemble
                    para.paragraph_format.keep_with_next = True
                    para.paragraph_format.keep_together = True
    except Exception as e:
        # En cas d'erreur, logger et ajouter un message d'erreur dans le document
        import traceback
        error_msg = f"Error rendering service: {str(e)}"
        print(f"ERROR in _render_service: {error_msg}")
        print(traceback.format_exc())
        # Ajouter un paragraphe d'erreur dans le document pour ne pas bloquer l'export
        try:
            p_err = doc.add_paragraph()
            r_err = p_err.add_run(f"[Error rendering service: {cat if 'cat' in locals() else 'unknown'}]")
            r_err.font.name = "Arial"
            r_err.font.size = Pt(NORMAL_PT)
            r_err.font.color.rgb = RGBColor(255, 0, 0)  # Rouge pour l'erreur
            _set_par_spacing(p_err, before_pt=SP_BEFORE_ALL_PT, after_pt=SP_AFTER_NORMAL_PT)
        except Exception:
            pass  # Si même l'ajout d'erreur échoue, on ignore

def _render_day_services(doc: Document, day, usable_w_cm: float = None) -> None:
    """Rend tous les services visibles (inclut Flight, Train, Ferry, Car Rental, etc.)."""
    lines = list(_get_attr(day, "lines", []) or [])
    if not lines:
        _blank_line(doc)
        return

    # Trier les lignes par position pour respecter l'ordre visuel
    lines = sorted(lines, key=lambda l: _get_attr(l, "position") or 0)

    # Filtrer les lignes pour exclure les Internal info
    visible_lines = []
    for line in lines:
        cat = (_get_attr(line, "category") or "").strip()
        if not cat.lower().startswith("internal"):
            visible_lines.append(line)
    
    # Rendre chaque service en vérifiant si le suivant est aussi un transport ou un hotel
    for i, line in enumerate(visible_lines):
        current_cat = (_get_attr(line, "category") or "").strip()
        is_current_transport = _is_transport_service(current_cat)
        
        # Vérifier si le service suivant est aussi un transport ou un hotel (pour Private Transfer)
        skip_blank = False
        if i + 1 < len(visible_lines):
            next_line = visible_lines[i + 1]
            next_cat = (_get_attr(next_line, "category") or "").strip()
            
            # Cas 1: Service de transport suivi d'un autre transport
            if is_current_transport and _is_transport_service(next_cat):
                skip_blank = True
            # Cas 2: Private Transfer suivi d'un Hotel
            elif current_cat == "Private Transfer" and next_cat in ("Hotel", "New Hotel"):
                skip_blank = True
        
        _render_service(doc, line, usable_w_cm or 0, skip_blank_line_after=skip_blank)

def _insert_day_block(doc: Document, qout, day, day_idx, usable_width_cm: float):
    # Jamais d'image hero pour le premier jour
    photo_para = None
    blank_para = None
    if day_idx > 0:
        decorative_images = _get_attr(day, "decorative_images", []) or []
        if len(decorative_images) >= 2:
            imgs = decorative_images[:2]
            # Garder un espace entre la photo et la date, mais s'assurer qu'ils ne sont pas séparés
            photo_para, blank_para = _insert_two_heroes(doc, imgs[0], imgs[1], add_blank_after=True)

    # Ligne de date
    date_iso = _get_attr(day, "date")
    base = fmt_ui_date(date_iso) if date_iso else ""
    if _is_first_of_dest_block(qout.days, day_idx):
        n = _count_nights_from_index(qout.days, day_idx)
        dest = (_get_attr(day, "destination") or "").strip()
        if dest and n > 0:
            base = f"{base} {dest} for {n} night{'s' if n > 1 else ''}"
    date_para = _add_date_line(doc, base)
    
    # S'assurer que la photo, la ligne vide et la date restent ensemble (pas de séparation)
    if photo_para is not None:
        photo_para.paragraph_format.keep_with_next = True
        if blank_para is not None:
            blank_para.paragraph_format.keep_with_next = True
        # S'assurer aussi que la date reste avec la photo
        date_para.paragraph_format.keep_with_next = True

    # Utiliser _render_day_services pour rendre tous les services (inclut Flight, Train, etc.)
    _render_day_services(doc, day, usable_width_cm)

def _apply_terms_conditions_style(para: Paragraph):
    """
    Applique le style spécifique aux termes et conditions :
    - Police Arial 8
    - Alignement : Gauche
    - Espacement avant : 0 pt
    - Espacement après : 0 pt
    - Interligne : Multiple 1.08
    - Mise en retrait : 0" avant et après
    - Gestion des veuves et des orphelins activée
    """
    # Alignement à gauche
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Formatage du paragraphe
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = None
    pf.right_indent = None
    pf.first_line_indent = None
    
    # Interligne Multiple 1.08
    # python-docx utilise line_spacing_rule et line_spacing pour le multiple
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.08
    
    # Gestion des veuves et des orphelins (widow/orphan control)
    para._element.get_or_add_pPr().set(qn('w:widowControl'), '1')

def _find_terms_heading_index(doc: Document) -> Optional[int]:
    """
    Find the index of the 'Essential Travel Terms and Conditions' heading.
    Returns paragraph index or None if not found.
    """
    target = "essential travel terms and conditions"
    for i, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip().lower() == target:
            return i
    return None

def _compute_usable_width_cm(section) -> float:
    page_cm = section.page_width.cm
    left_cm = section.left_margin.cm
    right_cm = section.right_margin.cm
    return max(0.0, page_cm - left_cm - right_cm)

def build_docx_for_quote(db: Session, quote_id: int) -> BytesIO:
    """
    Build a .docx for the given quote id, using the repository template and inserting
    generated content before the Terms & Conditions heading.
    """
    # Import here to avoid circular import
    from ...api.quotes import _to_out
    
    # Re-query properly:
    q = db.query(Quote).filter(Quote.id == quote_id).first()
    if not q:
        raise ValueError("Quote not found")
    qout = _to_out(q, db=db, include_first_image=True)

    # Load template
    doc = Document(str(WORD_TEMPLATE_PATH))
    section = doc.sections[0]
    usable_width_cm = _compute_usable_width_cm(section)

    # Find T&C anchor paragraph
    tc_idx = _find_terms_heading_index(doc)
    if tc_idx is None:
        # No T&C found: append content at end
        # Titre global
        title = (_get_attr(qout, "display_title") or _get_attr(qout, "title") or "").strip()
        if title:
            _add_title(doc, title)
        # Global heroes
        global_heroes = [_get_attr(qout, "hero_photo_1"), _get_attr(qout, "hero_photo_2")]
        global_heroes = [u for u in global_heroes if u]
        if len(global_heroes) >= 2:
            _insert_two_heroes(doc, global_heroes[0], global_heroes[1])  # Ignore returned tuple for global photos
        # Days - trier par position pour respecter l'ordre visuel
        days = sorted(qout.days or [], key=lambda d: _get_attr(d, "position") or 0)
        for i, day in enumerate(days):
            _insert_day_block(doc, qout, day, i, usable_width_cm)
    else:
        # Start a fresh doc with same section settings
        doc = Document()
        new_sec = doc.sections[0]
        old_sec = Document(str(WORD_TEMPLATE_PATH)).sections[0]
        new_sec.page_width = old_sec.page_width
        new_sec.page_height = old_sec.page_height
        new_sec.left_margin = old_sec.left_margin
        new_sec.right_margin = old_sec.right_margin
        new_sec.top_margin = old_sec.top_margin
        new_sec.bottom_margin = old_sec.bottom_margin
        usable_width_cm = _compute_usable_width_cm(new_sec)
        # Generate content area now on fresh doc
        title = (_get_attr(qout, "display_title") or _get_attr(qout, "title") or "").strip()
        if title:
            _add_title(doc, title)
        # Global heroes
        global_heroes = [_get_attr(qout, "hero_photo_1"), _get_attr(qout, "hero_photo_2")]
        global_heroes = [u for u in global_heroes if u]
        if len(global_heroes) >= 2:
            _insert_two_heroes(doc, global_heroes[0], global_heroes[1])  # Ignore returned tuple for global photos
        # Days - trier par position pour respecter l'ordre visuel
        days = sorted(qout.days or [], key=lambda d: _get_attr(d, "position") or 0)
        for i, day in enumerate(days):
            _insert_day_block(doc, qout, day, i, usable_width_cm)
        # Page break then re-add T&C with specific style
        doc.add_page_break()
        # Recharger le template pour récupérer le texte des termes et conditions
        template_doc = Document(str(WORD_TEMPLATE_PATH))
        template_tc_paragraphs = template_doc.paragraphs[tc_idx:]
        # Copier le texte et appliquer le style spécifique
        for source_para in template_tc_paragraphs:
            target_para = doc.add_paragraph()
            # Copier le texte en préservant les runs (pour garder gras/italique si nécessaire)
            if source_para.runs:
                for source_run in source_para.runs:
                    target_run = target_para.add_run(source_run.text)
                    # Préserver le formatage de base (gras, italique) mais appliquer Arial 8
                    target_run.bold = source_run.bold
                    target_run.italic = source_run.italic
                    target_run.underline = source_run.underline
                    target_run.font.name = 'Arial'
                    target_run.font.size = Pt(8)
                    # Préserver la couleur si elle existe
                    if source_run.font.color and source_run.font.color.rgb:
                        target_run.font.color.rgb = source_run.font.color.rgb
            else:
                # Paragraphe vide, ajouter un run vide avec Arial 8
                target_run = target_para.add_run("")
                target_run.font.name = 'Arial'
                target_run.font.size = Pt(8)
            # Appliquer le style de paragraphe spécifique
            _apply_terms_conditions_style(target_para)

    # Save to memory
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out
